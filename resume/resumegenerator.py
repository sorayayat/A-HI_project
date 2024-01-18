from docx import Document
import os
import subprocess

current_directory = os.getcwd()

current_directory = os.getcwd()


def format_skills(skills):
    formatted_skills = ''
    for i, skill in enumerate(skills):
        if i % 5 == 0 and i != 0:
            formatted_skills += '\n'  # 다섯 개마다 줄 바꿈 추가
        elif i != 0:
            formatted_skills += '  '  # 기술 사이에 두 개의 공백 추가
        formatted_skills += skill
    return formatted_skills

# 이력서 데이터를 받아서 텍스트 치환에 사용할 context를 생성


def generate_resume(resume_data):

    formatted_skills = format_skills(resume_data.get("skills", []))

    # Placeholder에 대해 초기 공백 값 설정
    context = {f'{{Experiences{i}}}': ' ' for i in range(1, 6)}
    context.update({f'{{ExperiencesDetail{i}}}': ' ' for i in range(1, 6)})
    context.update({f'{{Projects{i}}}': ' ' for i in range(1, 6)})
    context.update({f'{{ProjectsDetail{i}}}': ' ' for i in range(1, 6)})

    # 기존 로직을 사용하여 실제 데이터가 있는 경우 값 업데이트
    for i, experience in enumerate(resume_data.get("experiences", [])):
        context[f'{{Experiences{i + 1}}}'] = experience
    for i, detail in enumerate(resume_data.get("experiencesdetail", [])):
        context[f'{{ExperiencesDetail{i + 1}}}'] = detail
    for i, project in enumerate(resume_data.get("projects", [])):
        context[f'{{Projects{i + 1}}}'] = project
    for i, detail in enumerate(resume_data.get("projectsdetail", [])):
        context[f'{{ProjectsDetail{i + 1}}}'] = detail

    # 나머지 데이터는 기존 로직대로 처리
    context.update({
        '{Name}': resume_data.get("name", ' '),
        '{Phone}': resume_data.get("phonenumber", ' '),
        '{Email}': resume_data.get("email", ' '),
        '{Git}': resume_data.get("git", ' '),
        '{JobTitle}': resume_data.get("jobtitle", ' '),
        '{Skills}': formatted_skills,
        '{Experiences}': ', '.join(resume_data.get("experiences", [])),
        '{ExperiencesDetail}': '\n'.join(resume_data.get("experiencesdetail", [])),
        '{Projects}': '  '.join(resume_data.get("projects", [])),
        '{ProjectsDetail}': '\n'.join(resume_data.get("projectsdetail", [])),
        '{Education}': resume_data.get("education", ' '),
        '{EducationDetail}': resume_data.get("educationdetail", ' '),
        '{AwardsAndCertifications}': '\n'.join(resume_data.get("awardsandcertifications", []))
    })

    return context

# 텍스트 치환 함수


def replace_text_in_paragraph(paragraph, context):
    for run in paragraph.runs:
        for key, value in context.items():
            if key in run.text:
                run.text = run.text.replace(
                    key, ' ' if value is None or value == '' else value)

# 이력서 템플릿을 채우고 저장하는 함수


def fill_template(template_path, output_path, context):
    try:
        doc = Document(template_path)

        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, context)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, context)

        doc.save(output_path)
        print("Document created successfully. File path:", output_path)
    except Exception as e:
        print(f"Error: {e}")


# LibreOffice 실행 파일 경로 설정
libreoffice_path = 'C:\\Program Files\\LibreOffice\\program\\soffice.exe'

# Docx 파일을 PDF로 변환하는 함수


def convert_to_pdf(output_path_docx, output_path_pdf, libreoffice_path):
    try:
        subprocess.run([libreoffice_path, '--headless', '--convert-to', 'pdf',
                       output_path_docx, '--outdir', os.path.dirname(output_path_pdf)])
        print(f"PDF conversion successful. File path: {output_path_pdf}")
    except Exception as e:
        print(f"Error during PDF conversion: {e}")


def generate_resume_content(resume_data):

    user_name = resume_data.get("name", "Unnamed").replace(' ', '_')

    output_folder = os.path.join(current_directory, 'static/resumeResult')

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

     # 파일 경로에 사용자 이름 포함
    output_path_docx = os.path.join(
        output_folder, '{}_Resume.docx'.format(user_name))
    output_path_pdf = os.path.join(
        output_folder, '{}_Resume.pdf'.format(user_name))

    # resume_data를 확인해 적절한 템플릿 파일 선택
    experiences = resume_data.get("experiences")
    certifications = resume_data.get("awardsandcertifications")

    if experiences and certifications:
        template_path = os.path.join(
            current_directory, 'resume/tem', 'Template100.docx')  # 경력과 자격증 둘 다 존재하는 경우
    elif not experiences and certifications:
        template_path = os.path.join(
            current_directory, 'resume/tem', 'Template2.docx')  # 자격증만 있는 신입 템플릿
    elif experiences and not certifications:
        template_path = os.path.join(
            current_directory, 'resume/tem', 'Template3.docx')  # 경력만 있는 경력자 템플릿
    else:
        template_path = os.path.join(
            current_directory, 'resume/tem', 'Template4.docx')  # 경력과 자격증 모두 없는 신입 템플릿

    # 선택된 템플릿으로 context를 적용
    context = generate_resume(resume_data)

    # 이력서 템플릿을 채우고 저장
    fill_template(template_path, output_path_docx, context)

    # Docx 파일을 PDF로 변환
    convert_to_pdf(output_path_docx, output_path_pdf, libreoffice_path)

    return output_path_pdf  # PDF 파일 경로 반환
